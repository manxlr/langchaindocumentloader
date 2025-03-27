import os
import json
import faiss
import numpy as np
import pandas as pd
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.retrievers import BM25Retriever
from langchain.schema import Document
from docx import Document as DocxDocument
import xlrd

# Define Paths
DATA_DIRECTORY = r"data"
EMBEDDING_MODEL_PATH = r"model\bge_large"
FAISS_DB_PATH = r"faiss_index"
METADATA_PATH = r"faiss_metadata.json"

# Step 1: Function to Load All Supported File Types
def load_files_from_directory(directory):
    all_documents = []
    
    for root, _, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            
            if file.endswith(".pdf"):
                print(f"[INFO] Loading PDF: {file}")
                loader = PyPDFLoader(file_path)
                all_documents.extend(loader.load())

            elif file.endswith(".txt"):
                print(f"[INFO] Loading TXT: {file}")
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                all_documents.append(Document(page_content=text))

            elif file.endswith(".docx"):
                print(f"[INFO] Loading DOCX: {file}")
                doc = DocxDocument(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
                all_documents.append(Document(page_content=text))

            elif file.endswith(".csv"):
                print(f"[INFO] Loading CSV: {file}")
                df = pd.read_csv(file_path)
                text = df.to_string()
                all_documents.append(Document(page_content=text))

            elif file.endswith(".xls") or file.endswith(".xlsx"):
                print(f"[INFO] Loading EXCEL: {file}")
                df = pd.read_excel(file_path)
                text = df.to_string()
                all_documents.append(Document(page_content=text))

    return all_documents

# Load all files from the data directory
print("[INFO] Loading documents from multiple file types...")
docs = load_files_from_directory(DATA_DIRECTORY)
print(f"[INFO] Loaded {len(docs)} documents.")

# Step 2: Split Text into Chunks
print("[INFO] Splitting documents into large chunks...")
text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=750)
chunks = text_splitter.split_documents(docs)
print(f"[INFO] Split into {len(chunks)} chunks.")

# Step 3: Load Embedding Model
print("[INFO] Loading local high-quality embedding model...")
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_PATH)
print("[INFO] Embedding model loaded successfully.")

# Step 4: Generate Embeddings
print("[INFO] Generating embeddings...")
chunk_texts = [chunk.page_content for chunk in chunks]
chunk_embeddings = embeddings.embed_documents(chunk_texts)
embedding_array = np.array(chunk_embeddings).astype('float32')

# Step 5: Store in FAISS
print("[INFO] Creating FAISS vector database...")
index = faiss.IndexFlatL2(embedding_array.shape[1])
index.add(embedding_array)
faiss.write_index(index, FAISS_DB_PATH)
print(f"[INFO] FAISS index stored at: {FAISS_DB_PATH}")

# Step 6: Store Metadata Separately
print("[INFO] Storing metadata...")
metadata = [{"id": i, "text": chunks[i].page_content} for i in range(len(chunks))]
with open(METADATA_PATH, "w", encoding="utf-8") as f:
    json.dump(metadata, f, indent=4)
print(f"[INFO] Metadata stored at: {METADATA_PATH}")

# Step 7: Hybrid Retrieval (FAISS + BM25)
bm25_retriever = BM25Retriever.from_documents([Document(page_content=chunk.page_content) for chunk in chunks])

class HybridRetriever:
    def __init__(self, dense_index, sparse_retriever, metadata):
        self.dense_index = dense_index
        self.sparse_retriever = sparse_retriever
        self.metadata = metadata

    def invoke(self, query, top_k=3):
        # 1. Dense retrieval
        query_embedding = np.array(embeddings.embed_query(query)).astype('float32').reshape(1, -1)
        distances, indices = self.dense_index.search(query_embedding, top_k)
        # Convert dense results to dict with {'text': ...}
        dense_results = [{'text': self.metadata[i]['text']} for i in indices[0]]

        # 2. Sparse (BM25) retrieval → list of Documents
        bm25_docs = self.sparse_retriever.invoke(query)

        # 3. Convert BM25 Documents to dict with {'text': ...}
        sparse_results = [{'text': doc.page_content} for doc in bm25_docs]

        # 4. Combine and return
        return dense_results[:2] + sparse_results[:2]


# Load FAISS Index & Metadata
index = faiss.read_index(FAISS_DB_PATH)
with open(METADATA_PATH, "r", encoding="utf-8") as f:
    metadata = json.load(f)

retriever = HybridRetriever(index, bm25_retriever, metadata)

# Step 8: Test Search
query = "What is the role of gender in The Left Hand of Darkness"
#query = "What are the key themes in these documents?"
results = retriever.invoke(query)

print("\n[INFO] Top Retrieved Documents:")
for i, result in enumerate(results, 1):
    print(f"\nResult {i}:\n{result['text'][:500]}")

print("\n[INFO] FAISS Processing Completed Successfully!")
